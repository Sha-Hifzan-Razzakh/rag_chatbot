from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    doc_metadata: Dict[str, Any]
    pages: List[Dict[str, Any]]  # for PDFs we keep per-page info; otherwise empty


def _read_text_file(path: Path) -> str:
    """
    Read .txt / .md safely with fallback encoding.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # fallback; not perfect but avoids hard crashes for many files
        return path.read_text(encoding="latin-1", errors="ignore")


def _load_docx(path: Path) -> str:
    """
    Extract text from a DOCX using python-docx.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "DOCX support requires 'python-docx'. Please add it to backend requirements."
        ) from e

    doc = Document(str(path))
    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # tables (optional but usually helpful)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    return "\n\n".join(parts).strip()


def _infer_source_type(ext: str) -> str:
    ext = ext.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".md":
        return "markdown"
    return "text"


def extract_document(
    file_path: Path,
    *,
    original_filename: Optional[str] = None,
    content_type: Optional[str] = None,
    enable_ocr: bool = False,
    ocr_language: str = "eng",
    ocr_dpi: int = 200,
    ocr_min_chars: int = 50,
    max_pages: Optional[int] = None,
) -> ExtractionResult:
    """
    Route a file to the appropriate extractor and return normalized output.

    Args:
        file_path: path to saved upload on disk
        original_filename: user-provided filename
        content_type: MIME type (optional)
        enable_ocr: OCR fallback for scanned PDFs
        ocr_language / ocr_dpi / ocr_min_chars: OCR settings
        max_pages: limit pages for PDFs (useful for huge documents)

    Returns:
        ExtractionResult(text, doc_metadata, pages)
    """
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()
    source_type = _infer_source_type(ext)

    doc_metadata: Dict[str, Any] = {
        "file_name": file_path.name,
        "original_filename": original_filename or file_path.name,
        "content_type": content_type,
        "source_type": source_type,
        "source_path": str(file_path),
    }

    # ---- PDF ----
    if ext == ".pdf":
        from .loaders.pdf_loader import load_pdf

        text, pdf_meta, pages = load_pdf(
            file_path,
            max_pages=max_pages,
            enable_ocr=enable_ocr,
            ocr_language=ocr_language,
            ocr_dpi=ocr_dpi,
            ocr_min_chars=ocr_min_chars,
        )
        # merge, but keep extractor fields too
        doc_metadata.update(pdf_meta)
        return ExtractionResult(text=text, doc_metadata=doc_metadata, pages=pages)

    # ---- DOCX ----
    if ext == ".docx":
        text = _load_docx(file_path)
        doc_metadata["num_pages"] = None
        doc_metadata["char_count"] = len(text)
        return ExtractionResult(text=text, doc_metadata=doc_metadata, pages=[])

    # ---- TXT / MD (and other plain text-like) ----
    if ext in {".txt", ".md"}:
        text = _read_text_file(file_path).strip()
        doc_metadata["num_pages"] = None
        doc_metadata["char_count"] = len(text)
        return ExtractionResult(text=text, doc_metadata=doc_metadata, pages=[])

    # If validators allowed it, you likely won't get here; still be safe:
    raise ValueError(f"Unsupported file type: '{ext}' for file '{file_path.name}'")


def extract_many(
    file_paths: List[Path],
    *,
    originals: Optional[List[str]] = None,
    content_types: Optional[List[str]] = None,
    enable_ocr: bool = False,
    ocr_language: str = "eng",
    ocr_dpi: int = 200,
    ocr_min_chars: int = 50,
    max_pages: Optional[int] = None,
) -> List[ExtractionResult]:
    """
    Convenience helper for batch extraction.
    """
    results: List[ExtractionResult] = []
    for idx, p in enumerate(file_paths):
        results.append(
            extract_document(
                p,
                original_filename=(originals[idx] if originals and idx < len(originals) else None),
                content_type=(content_types[idx] if content_types and idx < len(content_types) else None),
                enable_ocr=enable_ocr,
                ocr_language=ocr_language,
                ocr_dpi=ocr_dpi,
                ocr_min_chars=ocr_min_chars,
                max_pages=max_pages,
            )
        )
    return results
